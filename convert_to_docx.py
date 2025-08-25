import markdown
from docx import Document
import sys
from docx import Document
import markdown
from bs4 import BeautifulSoup


def convert_md_to_docx(md_file, docx_file, base_docx="Context/Profil_v.-Mengershausen-AI-Process-Opt.docx"):
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    html = markdown.markdown(md_content)
    soup = BeautifulSoup(html, 'html.parser')

    document = Document(base_docx)
    for element in soup.children:
        if element.name == 'h1':
            document.add_heading(element.text, level=1)
        elif element.name == 'h2':
            document.add_heading(element.text, level=2)
        elif element.name == 'h3':
            document.add_heading(element.text, level=3)
        elif element.name == 'p':
            document.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'hr':
            document.add_page_break()

    document.save(docx_file)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_to_docx.py <input_md_file> <output_docx_file>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    convert_md_to_docx(input_file, output_file)
