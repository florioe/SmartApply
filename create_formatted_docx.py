import docx
import sys

def create_formatted_docx(md_file, reference_docx, output_docx):
    """
    Converts a markdown file to a formatted DOCX file,
    using styles from a reference DOCX.
    """
    # Use the reference document to inherit its styles
    document = docx.Document(reference_docx)

    # Clear all existing content (paragraphs and tables) from the reference doc
    # to start with a blank page that has the correct styles.
    for para in document.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    for table in document.tables:
        t = table._element
        t.getparent().remove(t)

    with open(md_file, 'r', encoding='utf-8') as f:
        for line in f:
            stripped_line = line.strip()
            if not stripped_line:
                # Add an empty paragraph for blank lines to preserve spacing
                document.add_paragraph()
                continue

            if stripped_line.startswith('# '):
                # Use 'Heading 1' style for the main heading
                p = document.add_paragraph(stripped_line[2:], style='Heading 1')
            elif stripped_line.startswith('## '):
                # Use 'Heading 2' style for subheadings
                p = document.add_paragraph(stripped_line[3:], style='Heading 2')
            elif stripped_line.startswith('Rolle:'):
                # Use 'Heading 3' style for roles in project history
                p = document.add_paragraph(stripped_line, style='Heading 3')
            elif stripped_line.startswith('* '):
                # Use 'List Bullet' style for bullet points
                p = document.add_paragraph(stripped_line[2:], style='List Bullet')
            elif stripped_line == '\\newpage':
                document.add_page_break()
            else:
                # Use 'Normal' style for all other text
                p = document.add_paragraph(stripped_line)

    document.save(output_docx)

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python create_formatted_docx.py <input_md_file> <reference_docx_file> <output_docx_file>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    reference_file = sys.argv[2]
    output_file = sys.argv[3]
    
    create_formatted_docx(input_file, reference_file, output_file)
